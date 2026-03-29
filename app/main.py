# app/main.py
from fastapi import FastAPI, Depends, HTTPException, Query, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session
from typing import List, Optional
import requests
import re
from docx import Document
from app import auth
from app.database import get_db, init_db
from app.models import MethodicEntry, QAEntry
from app.search import (
    search_methodics_with_context,
    format_context_for_prompt,
    search_qa_entries,
    search_methodic_texts
)
from app.config import settings
from pydantic import BaseModel

app = FastAPI(title="Methodics Chat Bot (Dual Database)", version="3.1.0")
auth.init_auth(app)

# ------------------ CORS ------------------
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ------------------ DB INIT ------------------
@app.on_event("startup")
def on_startup():
    init_db()


# ------------------ MODELS ------------------
class ChatRequest(BaseModel):
    question: str
    max_results: int = 5
    full: bool = False


class MethodicSnippet(BaseModel):
    id: int
    title: str
    author: Optional[str] = None
    content_snippet: str


class ChatResponse(BaseModel):
    answer: str
    sources: List[MethodicSnippet]
    found_methodics: int

    class UploadResponse(BaseModel):
        id: int
        title: str
        author: str | None
        content_snippet: str

def detect_question_type(question: str) -> str:
    q = question.lower().strip()

    if q.startswith("что такое"):
        return "definition"

    if "какие методы" in q or q.startswith("какие методы"):
        return "methods"

    if "роль" in q:
        return "role"

    if "преимущества" in q:
        return "advantages"

    return "generic"


# ------------------ HELPERS ------------------
def fix_text(text: str) -> str:
    """
    Нормализует текст:
    - убирает лишние пробелы
    - пытается частично исправить склеенные фрагменты
    """
    if not text:
        return ""

    text = re.sub(r'\s+', ' ', text).strip()

    # частичная правка слепленных слов вида "текствузе"
    text = re.sub(r'([а-яё])([А-ЯA-Z])', r'\1 \2', text)

    # правка пробелов перед знаками препинания
    text = re.sub(r'\s+([.,!?;:])', r'\1', text)

    return text.strip()


def is_quality_answer(answer: str, question: str) -> bool:
    """
    Упрощенная проверка качества ответа от Gemini.
    Не режем хорошие ответы слишком агрессивно.
    """
    if not answer:
        return False

    answer = fix_text(answer)

    if len(answer.strip()) < 80:
        return False

    bad_markers = [
        "не могу ответить",
        "недостаточно информации",
        "контекст не содержит",
        "не найдено информации"
    ]
    answer_lower = answer.lower()
    if any(marker in answer_lower for marker in bad_markers):
        return False

    return True

def call_gemini_api(question: str, context: str) -> str:
    """
    Gemini используется как инструмент для краткого и понятного
    ответа строго по найденному контексту.
    """
    instruction = f"""
Ты — эксперт по педагогике высшего образования.

Ответь на вопрос, используя ТОЛЬКО предоставленный контекст.

Правила ответа:
1. Сначала кратко ответь на вопрос 1-2 предложениями.
2. Затем, если уместно, перечисли 2-4 ключевые особенности или вывода.
3. Не копируй контекст дословно большими кусками.
4. Не придумывай факты, которых нет в контексте.
5. Если информации недостаточно, так и напиши.
6. Пиши простым, понятным, академически нейтральным языком.

Вопрос:
{question}

Контекст:
{context[:5000]}

Ответ:
"""

    url = f"{settings.GEMINI_API_URL}?key={settings.GEMINI_API_KEY}"
    headers = {"Content-Type": "application/json"}
    body = {
        "contents": [{"parts": [{"text": instruction}]}],
        "generationConfig": {
            "temperature": 0.3,
            "maxOutputTokens": 700,
            "topP": 0.9,
            "topK": 40
        }
    }

    try:
        resp = requests.post(url, headers=headers, json=body, timeout=30)

        if resp.status_code != 200:
            print(f"Ошибка Gemini: {resp.status_code} - {resp.text[:300]}")
            return ""

        data = resp.json()
        answer = (
            data.get("candidates", [{}])[0]
            .get("content", {})
            .get("parts", [{}])[0]
            .get("text", "")
            .strip()
        )

        return fix_text(answer)

    except Exception as e:
        print(f"Ошибка обращения к Gemini: {e}")
        return ""


def synthesize_answer(search_results: dict, question: str) -> str:
    q_type = detect_question_type(question)
    contexts = search_results.get("methodic_contexts", [])

    if not contexts:
        return "В методических материалах не найдено информации по данному вопросу."

    sentences = []
    seen = set()

    # Собираем предложения из найденных методичек
    for ctx in contexts:
        for s in ctx.get("relevant_sentences", []):
            s = fix_text(s)
            key = s.lower()

            if len(s) >= 50 and key not in seen:
                sentences.append(s)
                seen.add(key)

    if not sentences:
        return "В методических материалах отсутствует содержательная информация по данному вопросу."

    # Берём до 3 лучших фрагментов
    top_sentences = sentences[:3]

    # -------- МЕТОДЫ --------
    if q_type == "methods":
        answer_parts = [
            "По найденным методическим материалам можно выделить следующие методы и подходы:"
        ]
        for i, sentence in enumerate(top_sentences, 1):
            answer_parts.append(f"{i}. {sentence}")
        return " ".join(answer_parts)

    # -------- ОПРЕДЕЛЕНИЕ --------
    if q_type == "definition":
        first = top_sentences[0]
        rest = top_sentences[1:3]

        answer_parts = [
            f"По методическим материалам суть понятия можно описать так: {first}"
        ]

        if rest:
            answer_parts.append("Дополнительно можно выделить следующее:")
            for i, sentence in enumerate(rest, 1):
                answer_parts.append(f"{i}. {sentence}")

        return " ".join(answer_parts)

    # -------- РОЛЬ --------
    if q_type == "role":
        answer_parts = [
            "По найденным материалам ключевая роль заключается в следующем:"
        ]
        for i, sentence in enumerate(top_sentences, 1):
            answer_parts.append(f"{i}. {sentence}")
        return " ".join(answer_parts)

    # -------- ПРЕИМУЩЕСТВА --------
    if q_type == "advantages":
        answer_parts = [
            "По найденным материалам можно выделить следующие преимущества:"
        ]
        for i, sentence in enumerate(top_sentences, 1):
            answer_parts.append(f"{i}. {sentence}")
        return " ".join(answer_parts)

    # -------- ОБЩИЙ СЛУЧАЙ --------
    answer_parts = [
        "По найденным методическим материалам можно выделить следующее:"
    ]
    for i, sentence in enumerate(top_sentences, 1):
        answer_parts.append(f"{i}. {sentence}")

    return " ".join(answer_parts)


def parse_methodic_docx(file) -> dict:
    """
    Парсинг .docx файла методички с таблицей.
    Возвращает dict с title, author, text.
    """
    doc = Document(file)

    # Если есть таблицы, берём первую таблицу
    if doc.tables:
        table = doc.tables[0]
        # Предполагаем, что первая строка таблицы — заголовки: Автор, Название, Текст
        # Вторая строка — сами данные
        if len(table.rows) < 2:
            raise ValueError("В таблице недостаточно строк (ожидается хотя бы заголовок и данные)")

        author = table.cell(1, 0).text.strip() or None
        title = table.cell(1, 1).text.strip() or "Без названия"
        text = table.cell(1, 2).text.strip() or ""

    else:
        # fallback на старый метод
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            raise ValueError("Пустой документ")
        author = paragraphs[0] if len(paragraphs) > 0 else None
        title = paragraphs[1] if len(paragraphs) > 1 else "Без названия"
        text = " ".join(paragraphs[2:]) if len(paragraphs) > 2 else ""

    return {
        "title": title,
        "author": author,
        "text": text
    }

# ------------------ MAIN LOGIC ------------------
@app.post("/chat", response_model=ChatResponse)
async def chat_with_methodics(
        request: ChatRequest,
        db: Session = Depends(get_db)
):
    """
    Улучшенный алгоритм:
    1. Сначала ищем похожие вопросы в qa_entries
    2. Если найдены - возвращаем готовые ответы
    3. Если нет - ищем в methodic_entries и обрабатываем через Gemini
    4. Проверяем качество ответа от Gemini
    """
    print(f" {'=' * 50}")
    print(f"Вопрос: {request.question}")

    # --- Шаг 1: Ищем в базе готовых Q&A ---
    qa_results = search_qa_entries(db, request.question, threshold=0.6, limit=request.max_results)

    if qa_results:
        print(f"Найдено {len(qa_results)} готовых ответов в Q&A")

        # Формируем ответ из Q&A
        if len(qa_results) == 1:
            answer = qa_results[0].answer
        else:
            # Объединяем несколько ответов
            answer_parts = ["Найдено несколько похожих вопросов:"]
            for i, qa in enumerate(qa_results[:3], 1):  # Берем до 3 ответов
                answer_parts.append(f" {i}. {qa.answer}")
            answer = " ".join(answer_parts)


        sources = []
        for qa in qa_results[:3]:
            if qa.methodic:
                sources.append(
                    MethodicSnippet(
                        id=qa.methodic.id,
                        title=qa.methodic.source_title or "Без названия",
                        author=qa.methodic.author,
                        content_snippet=f"Связанный вопрос: {qa.question[:150]}..."
                    )
                )

        return ChatResponse(
            answer=answer,
            sources=sources,
            found_methodics=len(qa_results)
        )

    # --- Шаг 2: Ищем в полных текстах методичек ---
    print("Q&A не найдены, ищем в полных текстах...")
    search_results = search_methodics_with_context(db, request.question, request.max_results)

    # Если ничего не найдено
    if not search_results['methodic_contexts']:
        print("Ничего не найдено в текстах методичек.")
        answer = (
            "По вашему запросу не найдено информации в методических материалах. "
            "Попробуйте переформулировать вопрос или обратитесь к администратору."
        )
        return ChatResponse(answer=answer, sources=[], found_methodics=0)

    print(f"Найдено {len(search_results['methodic_contexts'])} релевантных методичек")

    # --- Шаг 3: Формируем контекст и отправляем в Gemini ---
    context = format_context_for_prompt(search_results, request.question)  # Передаем вопрос
    print(f"Длина контекста: {len(context)} символов")

    gemini_answer = call_gemini_api(request.question, context)

    # --- Шаг 4: Проверяем качество ответа Gemini ---
    if gemini_answer and is_quality_answer(gemini_answer, request.question):
        print("Gemini дал качественный ответ")
        answer = gemini_answer
    else:
        print("Gemini не дал качественного ответа, формируем смысловой ответ")
        answer = synthesize_answer(search_results, request.question)

    # --- Шаг 5: Формируем источники для ответа ---
    sources = []
    if search_results['methodic_contexts']:
        for ctx in search_results['methodic_contexts'][:5]:
            methodic = ctx['methodic']

            # Формируем осмысленный сниппет
            if ctx['relevant_sentences']:
                # Берем самое релевантное предложение
                best_sentence = ctx['relevant_sentences'][0]
                clean_sentence = re.sub(r'\s+', ' ', best_sentence).strip()
                if len(clean_sentence) > 300:
                    clean_sentence = clean_sentence[:300] + "..."
                snippet = clean_sentence
            else:
                snippet = methodic.methodic_text[:200] + "..." if methodic.methodic_text else ""

            sources.append(
                MethodicSnippet(
                    id=methodic.id,
                    title=methodic.source_title or "Без названия",
                    author=methodic.author,
                    content_snippet=snippet
                )
            )

    print(f"Ответ сформирован, источников: {len(sources)}")

    return ChatResponse(
        answer=answer,
        sources=sources,
        found_methodics=len(search_results['methodic_contexts'])
    )


# ------------------ SEARCH ENDPOINT ------------------
@app.get("/search", response_model=List[MethodicSnippet])
async def search_methodics_endpoint(
        query: str = Query(..., description="Поисковый запрос"),
        limit: int = Query(10, description="Максимальное количество результатов"),
        db: Session = Depends(get_db)
):
    methodic_results = search_methodic_texts(db, query, limit)

    sources = []
    for methodic in methodic_results:
        preview = methodic.methodic_text[:200] + "..." if methodic.methodic_text and len(
            methodic.methodic_text) > 200 else methodic.methodic_text

        sources.append(
            MethodicSnippet(
                id=methodic.id,
                title=methodic.source_title or "Без названия",
                author=methodic.author,
                content_snippet=preview or ""
            )
        )

    return sources


# ------------------ GET METHODIC BY ID ------------------
@app.get("/methodics/{methodic_id}", response_model=MethodicSnippet)
async def get_methodic(methodic_id: int, db: Session = Depends(get_db)):
    methodic = db.query(MethodicEntry).filter(MethodicEntry.id == methodic_id).first()
    if not methodic:
        raise HTTPException(status_code=404, detail="Методичка не найдена")

    return MethodicSnippet(
        id=methodic.id,
        title=methodic.source_title or "Без названия",
        author=methodic.author,
        content_snippet=methodic.methodic_text or ""
    )
# ------------------ ADD METHODIC (DOCX UPLOAD) ------------------
@app.post("/methodics/upload", response_model=MethodicSnippet)
async def upload_methodic(
        file: UploadFile = File(...),
        db: Session = Depends(get_db)
):
    """
    ### Загрузка методического документа (.docx)

    Загружает, парсит и сохраняет новый методический материал в базу данных.
    Это ключевой эндпоинт для наполнения базы знаний.

    #### Ожидаемый формат входного файла (.docx)
    Документ **обязательно** должен содержать **первую таблицу**, которая используется для извлечения метаданных и текста.

    | Колонка 1 | Колонка 2 | Колонка 3 |
    |---|---|---|
    | **Автор** (Иванов А.А.) | **Название** (Методика анализа рисков) | **Текст** (Полный текст методики...) |

    *Примечание: Если таблица отсутствует, используется fallback-логика парсинга первых абзацев.*

    ---

    #### Параметры Запроса (Body)

    | Параметр | Тип | Описание | Обязательность |
    |---|---|---|---|
    | `file` | `binary` (.docx) | Файл методики. | **Да** |

    #### 200 OK - Тело Ответа

    Возвращает объект `MethodicSnippet`, представляющий сохраненную запись.

    ```json
    {
        "id": 15,
        "title": "Методика анализа рисков",
        "author": "Сидоров П.А.",
        "content_snippet": "Сохраненный текст методики, обрезанный до первых 300 символов..."
    }
    ```

    #### Возможные Ошибки (HTTPException)

    * **400 Bad Request**:
        * **Детали**: `"Можно загружать только .docx файлы"` (Неверный формат файла).
        * **Детали**: `"Ошибка чтения файла: [сообщение об ошибке]"` (Пустой документ, неверная структура таблицы, или другие ошибки парсинга).
    """


    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Можно загружать только .docx файлы")

    try:
        parsed = parse_methodic_docx(file.file)
        title = parsed["title"]
        author = parsed["author"]
        text = parsed["text"]
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Ошибка чтения файла: {str(e)}")

    # Сохраняем в базу
    methodic = MethodicEntry(
        source_title=title,
        author=author,
        methodic_text=text
    )
    db.add(methodic)
    db.commit()
    db.refresh(methodic)

    # Формируем сниппет
    snippet = text[:300] + "..." if len(text) > 300 else text

    return MethodicSnippet(
        id=methodic.id,
        title=title,
        author=author,
        content_snippet=snippet
    )


# ------------------ Q&A SEARCH ENDPOINT ------------------
@app.get("/qa/search")
async def search_qa(
        query: str = Query(..., description="Поисковый запрос"),
        threshold: float = Query(0.5, description="Порог схожести (0-1)"),
        limit: int = Query(5, description="Максимальное количество результатов"),
        db: Session = Depends(get_db)
):
    qa_results = search_qa_entries(db, query, threshold, limit)

    results = []
    for qa in qa_results:
        results.append({
            "id": qa.id,
            "question": qa.question,
            "answer": qa.answer,
            "methodic_title": qa.methodic.source_title if qa.methodic else None,
            "methodic_author": qa.methodic.author if qa.methodic else None
        })

    return {"results": results, "count": len(results)}


# ------------------ ROOT ENDPOINT ------------------
@app.get("/")
async def root():
    return {
        "message": "Methodics Chat Bot API (Dual Database)",
        "version": "3.1.0",
        "endpoints": [
            "POST /chat - Чат с поиском по Q&A и методичкам",
            "GET /search - Поиск по методичкам",
            "GET /qa/search - Поиск по Q&A",
            "GET /methodics/{id} - Получить методичку по ID"
        ]
    }


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)